import os
import pandas as pd
from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from langchain.schema import Document as LangChainDocument

def iter_block_items(parent):
    """
    顺序遍历 docx 中的段落和表格，保持上下文的连贯性。
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unrecognized parent element")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def extract_table_markdown(table, context_str="", chunk_size=300, overlap=50):
    """
    将表格转为 Markdown 格式。
    针对超大表格（超过 chunk_size），按行切片，但每一片都保留**表头**和**上下文描述**。
    这样保证了高准确率的 RAG 召回。
    """
    chunks = []
    if not table.rows:
        return chunks
        
    # 获取表头
    header_cells = [cell.text.replace('\n', ' ').strip() for cell in table.rows[0].cells]
    header_md = "| " + " | ".join(header_cells) + " |"
    separator_md = "|" + "|".join(["---"] * len(header_cells)) + "|"
    
    current_chunk_rows = []
    current_length = len(context_str) + len(header_md) + len(separator_md) + 2 # 换行符
    
    for i, row in enumerate(table.rows[1:], start=1):
        cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
        row_md = "| " + " | ".join(cells) + " |"
        
        # 简单按字符长度切分，实际可按 token，此处遵循约束 chunk_size 300
        if current_length + len(row_md) > chunk_size and current_chunk_rows:
            # 组装当前的 chunk
            table_md = "\n".join([header_md, separator_md] + current_chunk_rows)
            chunk_text = f"{context_str}\n{table_md}".strip()
            chunks.append(chunk_text)
            
            # overlap: 保留最后几行（粗略计算）
            overlap_rows = current_chunk_rows[-2:] if len(current_chunk_rows) >= 2 else current_chunk_rows
            current_chunk_rows = overlap_rows + [row_md]
            current_length = len(context_str) + len(header_md) + len(separator_md) + sum(len(r) for r in current_chunk_rows)
        else:
            current_chunk_rows.append(row_md)
            current_length += len(row_md)
            
    if current_chunk_rows:
        table_md = "\n".join([header_md, separator_md] + current_chunk_rows)
        chunk_text = f"{context_str}\n{table_md}".strip()
        chunks.append(chunk_text)
        
    return chunks

def load_docx_with_tables(file_path: str, chunk_size=300, overlap=50) -> list[LangChainDocument]:
    """
    加载并切分 docx，专门针对表格进行结构化抽取。
    返回 LangChain Document 列表。
    """
    doc = Document(file_path)
    documents = []
    
    current_text = ""
    last_paragraph = "" # 记录上一个段落，通常是表格的标题/描述
    
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                last_paragraph = text
                if len(current_text) + len(text) > chunk_size:
                    documents.append(LangChainDocument(page_content=current_text, metadata={"source": file_path, "type": "text"}))
                    # Overlap
                    current_text = current_text[-overlap:] + "\n" + text if len(current_text) > overlap else text
                else:
                    current_text += "\n" + text if current_text else text
        elif isinstance(block, Table):
            # 遇到表格，先把前面的文本存进去
            if current_text:
                documents.append(LangChainDocument(page_content=current_text, metadata={"source": file_path, "type": "text"}))
                current_text = ""
                
            # 处理表格，带上最近的上下文
            table_chunks = extract_table_markdown(block, context_str=last_paragraph, chunk_size=chunk_size, overlap=overlap)
            for tc in table_chunks:
                documents.append(LangChainDocument(page_content=tc, metadata={"source": file_path, "type": "table"}))
                
    if current_text:
        documents.append(LangChainDocument(page_content=current_text, metadata={"source": file_path, "type": "text"}))
        
    return documents

def load_excel(file_path: str, chunk_size=300, overlap=50) -> list[LangChainDocument]:
    """
    使用 pandas 读取 excel，并按行转换为 Markdown 存入。
    """
    documents = []
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        # 填补空缺
        df = df.fillna("")
        
        header_md = "| " + " | ".join([str(c) for c in df.columns]) + " |"
        separator_md = "|" + "|".join(["---"] * len(df.columns)) + "|"
        
        current_chunk_rows = []
        current_length = len(header_md) + len(separator_md) + len(sheet_name) + 10
        
        for _, row in df.iterrows():
            cells = [str(val).replace('\n', ' ') for val in row.values]
            row_md = "| " + " | ".join(cells) + " |"
            
            if current_length + len(row_md) > chunk_size and current_chunk_rows:
                table_md = "\n".join([header_md, separator_md] + current_chunk_rows)
                chunk_text = f"Sheet: {sheet_name}\n{table_md}"
                documents.append(LangChainDocument(page_content=chunk_text, metadata={"source": file_path, "type": "excel"}))
                
                overlap_rows = current_chunk_rows[-2:] if len(current_chunk_rows) >= 2 else current_chunk_rows
                current_chunk_rows = overlap_rows + [row_md]
                current_length = len(header_md) + len(separator_md) + sum(len(r) for r in current_chunk_rows)
            else:
                current_chunk_rows.append(row_md)
                current_length += len(row_md)
                
        if current_chunk_rows:
            table_md = "\n".join([header_md, separator_md] + current_chunk_rows)
            chunk_text = f"Sheet: {sheet_name}\n{table_md}"
            documents.append(LangChainDocument(page_content=chunk_text, metadata={"source": file_path, "type": "excel"}))
            
    return documents
