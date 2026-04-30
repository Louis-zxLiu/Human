import os
import docx
import uuid
from typing import List, Dict, Any

class DocxParser:
    """
    专门解析 Docx 的工具类，核心功能是：
    1. 提取普通段落文本。
    2. 提取表格，并将表格转为大模型极易理解的 Markdown 格式。
    3. 在表格前附加上下文（如表格标题），保证切片后不丢失语义。
    """
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = docx.Document(file_path)
        self.file_name = os.path.basename(file_path)
        # 生成一个基于文件名的固定 UUID，方便后续做增量更新（相同文件产生相同的 ID）
        self.file_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, self.file_name))

    def _table_to_markdown(self, table) -> str:
        """将 python-docx 的 table 对象转换为 Markdown 格式字符串"""
        if not table.rows:
            return ""
            
        md_lines = []
        # 处理表头 (第一行)
        header_cells = [cell.text.replace('\n', ' ').strip() for cell in table.rows[0].cells]
        md_lines.append("| " + " | ".join(header_cells) + " |")
        
        # 处理分隔符
        separator = ["---"] * len(header_cells)
        md_lines.append("| " + " | ".join(separator) + " |")
        
        # 处理数据行
        for row in table.rows[1:]:
            row_cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            # 跳过完全为空的行
            if any(row_cells):
                md_lines.append("| " + " | ".join(row_cells) + " |")
                
        return "\n".join(md_lines)

    def _table_to_unrolled_rows(self, table, context: str) -> List[str]:
        """将表格展开为基于行的独立语义块，极大提升结构化数据的 RAG 召回率"""
        if not table.rows or len(table.rows) < 2:
            return []
            
        header_cells = [cell.text.replace('\n', ' ').strip() for cell in table.rows[0].cells]
        
        row_texts = []
        for row in table.rows[1:]:
            row_cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            if any(row_cells):
                row_parts = []
                for i in range(min(len(header_cells), len(row_cells))):
                    if row_cells[i]:
                        row_parts.append(f"{header_cells[i]}: {row_cells[i]}")
                
                row_str = "；".join(row_parts)
                if context:
                    row_str = f"【相关上下文】: {context} \n【具体信息】: {row_str}"
                else:
                    row_str = f"【表格提取数据】: {row_str}"
                row_texts.append(row_str)
                
        return row_texts

    def parse(self) -> List[Dict[str, Any]]:
        """
        顺序解析文档，返回 Chunk 列表。
        每个 Chunk 包含: content, metadata(包含 file_id, type 等)
        """
        chunks = []
        current_context = "" # 用于保存表格前的段落作为上下文
        
        # docx 的 element 顺序遍历比较复杂，这里用一个简化的块级遍历
        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                # 这是一个段落
                for p in self.doc.paragraphs:
                    if p._element == element:
                        text = p.text.strip()
                        if text:
                            # 如果文本较短且不以标点结尾，可能是标题，暂存为上下文
                            if len(text) < 50 and not text.endswith(('。', '.', '！', '!', '？', '?')):
                                current_context = text
                            
                            chunks.append({
                                "content": text,
                                "metadata": {
                                    "file_id": self.file_id,
                                    "source": self.file_name,
                                    "type": "text"
                                }
                            })
                        break
            elif element.tag.endswith('tbl'):
                # 这是一个表格
                for table in self.doc.tables:
                    if table._element == element:
                        # 使用基于行的展开方式，解决表格格式导致 RAG 召回率低、产生幻觉的问题
                        row_chunks = self._table_to_unrolled_rows(table, current_context)
                        for row_text in row_chunks:
                            chunks.append({
                                "content": row_text,
                                "metadata": {
                                    "file_id": self.file_id,
                                    "source": self.file_name,
                                    "type": "table_row"
                                }
                            })
                        current_context = "" # 使用完后清空
                        break
                        
        return chunks
