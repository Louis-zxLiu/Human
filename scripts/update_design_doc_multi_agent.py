from __future__ import annotations

import json
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "产品总体设计文档.docx"


def main() -> None:
    report_path = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "reports" / "full1200_multi_agent_review.json"
    report = json.loads(report_path.read_text(encoding="utf-8"))
    doc = Document(DOC_PATH)

    score = f"{float(report['overall_score']):.2f}"
    case_count = int(report["case_count"])
    failure_count = int(report["failure_count"])
    report_json = _relative(report_path)
    report_md = report_json.replace(".json", ".md")

    _set_paragraph(doc, "文档版本：", "文档版本：V1.6")
    _set_paragraph(doc, "修订日期：", "修订日期：2026.6.20")
    _set_exact(
        doc,
        "当前使用版本：V1.5。本文以当前工程实现、本轮 Agent 工具循环改造和 reports 下可核验评测报告为依据。",
        f"当前使用版本：V1.6。本文以当前工程实现、本轮多 Agent 语义规划、最终审核 LLM 改造和 {report_json} 可核验评测报告为依据。",
    )
    _append_revision_row(doc.tables[0])

    _set_exact(
        doc,
        "系统采用 Agent 工具循环架构。用户问题先由 QueryPlanner 生成首步计划，pipeline 将计划转换为标准 ToolCall，ToolRunner 调用景区事实、DOCX 检索、行为 SQL 或路线规划工具，再由 AgentLoopController 基于 ToolObservation 决定继续调用工具、追问或最终回答。",
        "系统采用多 Agent 语义规划与工具循环架构。主 Agent 先由 QueryPlanner 生成首步计划，专业子 Agent 再分别完成景点事实语义规划、行为分析语义规划或路线规划；确定性执行器只负责白名单工具调用、结构化事实读取和 SQL 编译执行；最终审核 LLM 只做审查，不直接改写答案。",
    )
    _set_contains(
        doc,
        "-> Agent 工具循环层：planner、ToolRunner、AgentLoopController、structured_fact、hybrid_rag、behavior_sql、route",
        "游客 / 管理员\n"
        "  -> 多模态交互层：React 前台、管理后台、语音输入、数字人视频、conversation_context\n"
        "  -> 服务编排层：FastAPI、短期会话记忆、日志、ASR/TTS、数字人生成、弱 GPS 会话\n"
        "  -> Agent 工具循环层：主 Agent、专业语义规划子 Agent、ToolRunner、AgentLoopController、answer_review\n"
        "  -> 确定性执行层：structured_fact、hybrid_rag、behavior_sql、route_planner、response_contract\n"
        "  -> 数据与证据层：SQLite 事实表、游客行为库、ChromaDB、DOCX 知识片段、评测报告",
    )
    _set_exact(
        doc,
        "传统 RAG 直接检索并生成答案，容易把统计问题、事实问题、闲聊和越界请求混在一起。本系统将 planner 定位为首步规划器，不依赖业务关键词硬分流；实际答复由 Agent 在工具 observation 之上进行二次决策。",
        "传统 RAG 直接检索并生成答案，容易把统计问题、事实问题、闲聊和越界请求混在一起。本系统将主 Agent 定位为编排器，不依赖业务关键词硬分流；FACT、ANALYTICS、ROUTE 等子 Agent 负责语义规划，工具层只执行经过白名单校验的确定性动作。",
    )
    _set_contains(
        doc,
        "执行时，工具返回的 ToolObservation 包含 status",
        "执行时，工具返回的 ToolObservation 包含 status、confidence、response_kind、evidence、refusal、missing_slots、suggested_next_tools 和 trace。pipeline 生成候选答案后交给最终审核 LLM；审核只返回 approved、issues、risk_level 和 repair_action，不生成替代答案。若审核未通过，repair_action 仅作为提示，主 Agent 重新选择并调用同类或其他专业子 Agent 修复，再进入下一轮审核。",
    )
    _set_exact(
        doc,
        "结构化景区事实层：负责景点位置、开放信息、建筑参数、文化属性等确定性内容。",
        "结构化景区事实层：负责景点位置、开放信息、建筑参数、文化属性等确定性内容；事实子 Agent 先做 LLM 语义规划，再由结构化执行器读取可信字段。",
    )
    _set_exact(
        doc,
        "语义 SQL 分析层：负责基于游客行为库的访问量、停留、消费、满意度、分组和 TopN 分析。",
        "语义 SQL 分析层：负责基于游客行为库的访问量、停留、消费、满意度、分组和 TopN 分析；分析子 Agent 只输出结构化查询计划，SQL 由确定性编译器生成。",
    )
    _set_exact(
        doc,
        "3. QueryPlanner 生成首步计划，包含 strategy、question_type、confidence、reasoning 和景区/景点槽位。",
        "3. QueryPlanner 作为主 Agent 首步规划器生成 strategy、question_type、confidence、reasoning 和景区/景点槽位。",
    )
    _set_exact(
        doc,
        "4. pipeline 将计划转换为 ToolCall，由 ToolRunner 调用 structured_fact、hybrid_rag、behavior_sql 或 route_planner。",
        "4. pipeline 将首步计划转换为 ToolCall，并由对应专业子 Agent 生成语义计划，ToolRunner 再调用 structured_fact、hybrid_rag、behavior_sql 或 route_planner。",
    )
    _set_exact(
        doc,
        "5. 工具返回 ToolObservation，包含 answer、evidence、refusal、warnings、status、confidence 和 trace。",
        "5. 确定性执行器返回 ToolObservation，包含 answer、evidence、refusal、warnings、status、confidence 和 trace。",
    )
    _set_exact(
        doc,
        "6. AgentLoopController 基于 observation 决定 `final_answer`、`call_tool` 或 `ask_clarification`。",
        "6. AgentLoopController 基于 observation 决定 `final_answer`、`call_tool` 或 `ask_clarification`；LLM 可用时不再静默退回启发式语义路由。",
    )
    _set_exact(
        doc,
        "7. pipeline 汇总工具观察结果，按统一 response contract 封装 `answer + rag_metadata`。",
        "7. pipeline 汇总工具观察结果后调用 answer_review；最终审核 LLM 只审查，不生成替代答案。",
    )
    _set_contains(
        doc,
        "审核通过后按统一 response contract",
        "8. 审核通过后按统一 response contract 封装 `answer + rag_metadata`；若审核不通过，repair_action 仅作为提示，主 Agent 重新选择并调用子 Agent 修复，最多进行受控修复循环。",
    )
    _insert_after(
        doc,
        "8. 审核通过后按统一 response contract 封装 `answer + rag_metadata`；若审核不通过，repair_action 仅作为提示，主 Agent 重新选择并调用子 Agent 修复，最多进行受控修复循环。",
        "9. log service 将 plan/evidence/refusal/warnings/observability 及 tools、agent_loop、answer_review、answer_review_repair trace 持久化到交互日志。",
    )
    _set_exact(
        doc,
        "4. 文本进入 Agent 工具循环 pipeline，生成答案、结构化元数据和工具 trace。",
        "4. 文本进入多 Agent 工具循环 pipeline，完成语义规划、确定性执行、最终审核和必要修复，生成答案、结构化元数据和工具 trace。",
    )
    _set_exact(
        doc,
        "2. ASR 转文本后，QueryPlanner 选择 structured_fact 作为首步工具。",
        "2. ASR 转文本后，QueryPlanner 选择 structured_fact 作为首步工具，事实子 Agent 识别景点和问题类型。",
    )
    _set_exact(
        doc,
        "3. ToolRunner 调用 structured_fact，查询结构化事实，返回 ToolObservation。",
        "3. ToolRunner 调用 structured_fact，确定性读取结构化事实字段，返回 ToolObservation。",
    )
    _set_exact(
        doc,
        "4. 若 observation 提示证据不足或问题需要深资料，AgentLoopController 可继续调用 hybrid_rag 补充检索片段。",
        "4. 若 observation 或最终审核 LLM 提示证据不足、表达不自然或风险较高，主 Agent 可继续调用 hybrid_rag 补充检索片段。",
    )
    _set_exact(
        doc,
        "5. pipeline 汇总工具结果，前台展示回答和必要的工具状态，数字人用短文本进行语音视频讲解。",
        "5. pipeline 在最终审核通过后返回回答和必要的工具状态，数字人用短文本进行语音视频讲解。",
    )
    _set_exact(
        doc,
        "项目中 `tests/unified_eval_cases.jsonl` 主评测集当前为 1200 行，覆盖结构化事实、DOCX RAG、行为 SQL、路线融合和边界拒答五类场景。当前可核验的最新报告为 `reports/unified_eval_agent_full1200.json/md`，层级为 full。",
        f"项目中 `tests/unified_eval_cases.jsonl` 主评测集当前为 {case_count} 行，覆盖结构化事实、DOCX RAG、行为 SQL、路线融合和边界拒答五类场景。当前可核验的最新报告为 `{report_json}` / `{report_md}`，层级为 full。",
    )
    _set_contains(
        doc,
        "本轮回归命令共运行",
        "除统一评测外，项目通过 `tests/test_tool_runner_agent_loop.py`、`tests/test_pipeline_general_chat.py`、`tests/test_router_cache_and_planner.py`、`tests/test_rag_contract.py`、`tests/test_answer_review_agent.py` 和 `tests/test_sql_semantic_rules.py` 验证 Agent 工具循环、闲聊保护、planner 路由、最终审核、语义 SQL 计划和统一响应契约。本轮回归命令共运行 40 个用例，结果为 OK。",
    )
    _set_exact(
        doc,
        "Agent 工具循环核心引擎：首步规划、工具执行、observation 自修正、闲聊保护和边界拒答。",
        "多 Agent 工具循环核心引擎：主 Agent 编排、专业子 Agent 语义规划、确定性工具执行、最终审核 LLM、审核失败修复和边界拒答。",
    )
    _set_exact(
        doc,
        "当前系统已经从“会回答问题的导览 Demo”升级为“面向真实场景的 Agent 工具循环导览服务系统”。其核心价值在于把游客交互、短期会话记忆、知识证据、行为数据、工具调用 trace、数字人讲解、后台运营和评测回放连接成一个可持续迭代的工程闭环。",
        "当前系统已经从“会回答问题的导览 Demo”升级为“面向真实场景的多 Agent 工具循环导览服务系统”。其核心价值在于把游客交互、短期会话记忆、知识证据、行为数据、工具调用 trace、最终审核、数字人讲解、后台运营和评测回放连接成一个可持续迭代的工程闭环。",
    )
    _set_exact(
        doc,
        "面向工程：执行前先规划，执行后看 observation，工具 trace 可追踪，统计优先 SQL，拒答有结构化原因。",
        "面向工程：主 Agent 先规划，子 Agent 做语义计划，执行器确定性落地，最终审核 LLM 只审查，工具 trace 和 repair trace 可追踪。",
    )
    _set_contains(
        doc,
        "面向评测：",
        f"面向评测：主评测集规模清晰，当前 full 评测总分 {score}/100，失败 {failure_count} 题，报告可核验并可复现更新。",
    )

    _update_summary_tables(doc, report, report_json, report_md)
    doc.save(DOC_PATH)


def _relative(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT).as_posix()
    except ValueError:
        return path.as_posix()


def _set_paragraph(doc: Document, prefix: str, text: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            paragraph.text = text
            return
    raise ValueError(f"paragraph prefix not found: {prefix}")


def _set_exact(doc: Document, old: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text == old:
            paragraph.text = new
            return
        if paragraph.text == new:
            return
    raise ValueError(f"paragraph not found: {old[:60]}")


def _set_contains(doc: Document, needle: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = new
            return
        if paragraph.text == new:
            return
    raise ValueError(f"paragraph containing text not found: {needle[:60]}")


def _insert_after(doc: Document, anchor: str, text: str) -> None:
    paragraphs = doc.paragraphs
    for paragraph in paragraphs:
        if paragraph.text == text:
            return
    for paragraph in paragraphs:
        if paragraph.text == anchor:
            new_paragraph = _insert_paragraph_after(paragraph, text)
            new_paragraph.style = paragraph.style
            return
    raise ValueError(f"anchor paragraph not found: {anchor[:60]}")


def _insert_paragraph_after(paragraph: Any, text: str = "") -> Any:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.clear()
    new_para.add_run(text)
    return new_para


def _append_revision_row(table: Any) -> None:
    for row in table.rows:
        if len(row.cells) > 1 and row.cells[1].text.strip() == "V1.6":
            row.cells[5].text = "时瑞宁"
            row.cells[6].text = "补充多 Agent 语义规划、最终审核 LLM 与最新 1200 题评测结果"
            return
    row = table.add_row()
    values = ["06", "V1.6", "2026.6.20", "朱文楦", "刘志轩", "时瑞宁", "补充多 Agent 语义规划、最终审核 LLM 与最新 1200 题评测结果"]
    for cell, value in zip(row.cells, values):
        cell.text = value


def _update_summary_tables(doc: Document, report: dict[str, Any], report_json: str, report_md: str) -> None:
    score = f"{float(report['overall_score']):.2f}"
    case_count = int(report["case_count"])
    failure_count = int(report["failure_count"])

    table1 = doc.tables[1]
    table1.rows[3].cells[1].text = "结构化事实、semantic SQL、hybrid RAG、route planner、最终审核 LLM、拒答边界"
    table1.rows[3].cells[2].text = "主 Agent 生成首步计划，专业子 Agent 输出语义计划，ToolRunner 以标准 schema 调用确定性执行器，answer_review 只审核；repair_action 仅作为主 Agent 修复提示"
    table1.rows[5].cells[1].text = f"主评测集 {case_count} 题，当前报告可回放"
    table1.rows[5].cells[2].text = f"tests 中主评测集为 {case_count} 行；当前 {report_json} 为 full 层级、{score}/100"

    table2 = doc.tables[2]
    table2.rows[3].cells[1].text = "主 Agent 首步规划、专业子 Agent 语义规划、ToolRunner 工具调用、Agent observation loop、最终审核 LLM、事实问答、semantic SQL、hybrid RAG、route planner、response contract"
    table2.rows[3].cells[2].text = "app/rag/planner.py、app/rag/fact_agent.py、app/rag/sql_agent.py、app/rag/answer_review_agent.py、app/rag/tool_runner.py、app/rag/agent_loop.py、app/rag/pipeline.py"

    table5 = doc.tables[5]
    if not any(row.cells[0].text.strip() == "answer_review" for row in table5.rows):
        row = table5.add_row()
        row.cells[0].text = "answer_review"
        row.cells[1].text = "最终回答上线前审核、字段泄漏检查、风险判断和修复动作选择"
        row.cells[2].text = "AnswerReviewAgent + pipeline repair loop"
        row.cells[3].text = "approved、issues、risk_level、repair_action、answer_review_repair"

    table9 = doc.tables[9]
    table9.rows[1].cells[1].text = f"{case_count} 题"
    table9.rows[2].cells[1].text = "full"
    table9.rows[2].cells[2].text = report_json
    table9.rows[3].cells[1].text = f"{case_count} 题"
    table9.rows[3].cells[2].text = f"{report_md}/json"
    table9.rows[4].cells[1].text = f"{score} / 100"
    table9.rows[4].cells[2].text = f"{report_md}/json"
    table9.rows[5].cells[1].text = str(failure_count)
    table9.rows[5].cells[2].text = f"{report_md}/json"

    table10 = doc.tables[10]
    stats = report.get("by_gold_source") or {}
    source_order = ["behavior_sql", "boundary", "docx_rag", "docx_structured", "fusion"]
    while len(table10.rows) < len(source_order) + 1:
        table10.add_row()
    for index, source in enumerate(source_order, start=1):
        row = table10.rows[index]
        source_stats = stats.get(source, {})
        row.cells[0].text = source
        row.cells[1].text = str(source_stats.get("count", ""))
        row.cells[2].text = str(source_stats.get("passed", ""))
        row.cells[3].text = f"{source_stats.get('accuracy', '')}"


if __name__ == "__main__":
    main()
