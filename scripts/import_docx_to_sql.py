import os
import sys
import sqlite3
import pandas as pd
import docx

# 添加项目根目录到 Python 路径，以解决找不到 app 模块的问题
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.core.config import resolve_path

def process_docx_to_sqlite():
    """
    Build the scenic fact layer from the Lingshan structured DOCX dataset.
    The output attractions table is the deterministic fact source for scenic
    Q&A and route explanation.
    """
    docx_path = resolve_path("data/knowledge_base/灵山胜境 景点结构化数据集.docx")
    db_path = resolve_path("data/processed/tourist_behavior.db")
    
    print(f"[*] 开始读取原始 Docx 数据: {docx_path}")
    if not os.path.exists(docx_path):
        print("[-] 错误: 找不到原始 Docx 文件！")
        return
        
    try:
        # 1. 使用 python-docx 提取表格数据
        doc = docx.Document(docx_path)
        if not doc.tables:
            print("[-] 错误: Docx 文件中没有找到表格！")
            return
            
        table = doc.tables[0]
        data = []
        keys = None
        for i, row in enumerate(table.rows):
            text = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            if i == 0:
                keys = text
            else:
                data.append(dict(zip(keys, text)))
                
        # 2. 转换为 DataFrame
        df = pd.DataFrame(data)
        
        # 3. 列名映射为英文
        column_mapping = {
            "景区名称": "scenic_name",
            "景点ID": "attraction_id",
            "景点名称": "attraction_name",
            "具体位置": "location",
            "建筑/景观参数": "architecture_params",
            "核心功能": "core_function",
            "文化内涵": "cultural_meaning",
            "详细介绍": "description",
            "游玩亮点": "highlights",
            "演艺/开放信息": "open_info",
            "备注": "remarks"
        }
        
        # 处理可能的列名包含空格等情况
        df.rename(columns=lambda x: x.strip().replace(' ', ''), inplace=True)
        # 这里特别处理一下 '文 化内涵' 这种奇葩的空格
        df.rename(columns={'文化内涵': 'cultural_meaning', '文 化内涵': 'cultural_meaning'}, inplace=True)
        df.rename(columns=column_mapping, inplace=True)
        
        # 确保目录存在
        os.makedirs(os.path.dirname(db_path), exist_ok=True)
        
        # 4. 写入 SQLite 数据库的 attractions 表
        print(f"[*] 正在将景点结构化数据写入 SQLite 数据库 (attractions 表): {db_path}")
        with sqlite3.connect(db_path) as conn:
            df.to_sql("attractions", conn, if_exists="replace", index=False)
            
        print("[+] 离线结构化景点数据转换完成！")
        print(f"    共处理了 {len(df)} 个景点记录。")
        
    except Exception as e:
        print(f"[-] 处理过程中发生错误: {e}")

if __name__ == "__main__":
    process_docx_to_sqlite()
